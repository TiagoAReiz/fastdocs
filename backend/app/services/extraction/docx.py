import io

import docx as python_docx

from app.services.extraction.base import ExtractionResult, Section


def extract(content: bytes) -> ExtractionResult:
    document = python_docx.Document(io.BytesIO(content))

    sections: list[Section] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush():
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        if text:
            sections.append(
                Section(
                    text=text,
                    metadata={"heading": current_heading} if current_heading else {},
                    kind="text",
                )
            )
        buffer.clear()

    for para in document.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            flush()
            current_heading = para.text.strip() or current_heading
            continue
        if para.text.strip():
            buffer.append(para.text)
    flush()

    return ExtractionResult(
        sections=sections,
        extraction_method="native",
        document_metadata={"paragraph_count": len(document.paragraphs)},
    )
